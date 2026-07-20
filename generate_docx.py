"""Génère DATABASE.docx à partir de la structure réelle de cisat_db."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = r"d:/Documents/base nations uni/web/cisat_project/DATABASE.docx"

# --------- helpers ---------------------------------------------------------

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(hdr_cells[i], '1F3A5F')
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# --------- document --------------------------------------------------------

doc = Document()

# Marges
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Police par défaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Titre
title = doc.add_heading('Documentation — Base de données cisat_db', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Projet CISAT — Système de collecte d\'indicateurs climatiques & environnementaux')
r.italic = True
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Snapshot au 2026-05-20')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph()

# ---------------------------------------------------------------- §1
add_heading(doc, '1. Informations générales', level=1)
add_table(doc,
    ['Élément', 'Valeur'],
    [
        ['SGBD',              'PostgreSQL 18'],
        ['Nom de la base',    'cisat_db'],
        ['Hôte / Port',       'localhost / 5432'],
        ['Utilisateur',       'postgres'],
        ['Encodage client',   'UTF-8'],
        ['Schéma principal',  'public'],
    ],
    col_widths=[5, 11])

doc.add_paragraph()
add_para(doc, 'Extensions installées', bold=True)
add_table(doc,
    ['Extension', 'Rôle'],
    [
        ['plpgsql',  'Langage procédural natif PostgreSQL (triggers, fonctions)'],
        ['unaccent', 'Recherche insensible aux accents (noms d\'indicateurs / pays)'],
    ],
    col_widths=[4, 12])

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Note : ')
r.bold = True
p.add_run("la base est gérée hors Django (managed = False dans les modèles). "
          "Le schéma est créé via SQL natif ; Django n'opère que les migrations techniques "
          "(extensions, auth, sessions).")

# ---------------------------------------------------------------- §2
doc.add_page_break()
add_heading(doc, '2. Vue d\'ensemble du schéma', level=1)
add_para(doc, 'La base se compose de deux groupes de tables.')

add_heading(doc, '2.1 Tables métier (5)', level=2)
add_table(doc,
    ['Table', 'Rôle', 'Lignes'],
    [
        ['pays',                   'Référentiel des pays',                             '1'],
        ['utilisateur',            'Comptes collecteurs / administrateurs',            '1'],
        ['indicateur',             'Catalogue des indicateurs CISAT & BASICSET',       '515'],
        ['metadonnee_indicateur',  'Métadonnées détaillées d\'un indicateur (1-1)',    '515'],
        ['donnee_collectee',       'Valeurs saisies par pays / année / indicateur',    '3'],
    ],
    col_widths=[5, 9, 2])

add_heading(doc, '2.2 Tables techniques Django (10)', level=2)
add_para(doc, "Gérées automatiquement par le framework, à ne pas modifier manuellement :")
p = doc.add_paragraph()
r = p.add_run('auth_group, auth_group_permissions, auth_permission, auth_user, '
              'auth_user_groups, auth_user_user_permissions, django_admin_log, '
              'django_content_type, django_migrations, django_session.')
r.font.size = Pt(10)
r.italic = True

# ---------------------------------------------------------------- §3
add_heading(doc, '3. Diagramme relationnel', level=1)
diagram = (
    "            ┌──────────┐\n"
    "            │   pays   │◄────────────┐\n"
    "            └────┬─────┘             │\n"
    "                 │                   │\n"
    "                 │ pays_id           │ pays_id\n"
    "                 ▼                   │\n"
    "         ┌──────────────┐            │\n"
    "         │ utilisateur  │            │\n"
    "         └──────┬───────┘            │\n"
    "                │                    │\n"
    "                │ utilisateur_id     │\n"
    "                │ valide_par         │\n"
    "                ▼                    │\n"
    "       ┌────────────────────┐        │\n"
    "       │  donnee_collectee  │────────┘\n"
    "       └─────────┬──────────┘\n"
    "                 │ indicateur_id\n"
    "                 ▼\n"
    "            ┌────────────┐           ┌──────────────────────────┐\n"
    "            │ indicateur │◄──────────│  metadonnee_indicateur   │ (1-1)\n"
    "            └────────────┘           └──────────────────────────┘\n"
)
p = doc.add_paragraph()
r = p.add_run(diagram)
r.font.name = 'Consolas'
r.font.size = Pt(9)

# ---------------------------------------------------------------- §4
doc.add_page_break()
add_heading(doc, '4. Détail des tables métier', level=1)

# --- pays
add_heading(doc, 'Table pays', level=2)
add_para(doc, 'Référentiel des pays participants.')
add_table(doc,
    ['Colonne', 'Type', 'Contraintes', 'Description'],
    [
        ['id',       'integer',      'PK, auto',          'Identifiant interne'],
        ['code_iso', 'character(3)', 'NOT NULL, UNIQUE',  'Code ISO 3166-1 alpha-3 (CIV, FRA, ...)'],
        ['nom',      'varchar(100)', 'NOT NULL',          'Nom officiel du pays'],
        ['region',   'varchar(100)', 'NULL',              'Région géographique'],
        ['actif',    'boolean',      'défaut true',       'Pays activé pour la collecte'],
    ],
    col_widths=[3, 3, 3, 7])
p = doc.add_paragraph()
p.add_run('Index : ').bold = True
p.add_run('PK sur id, UNIQUE sur code_iso.')
p = doc.add_paragraph()
p.add_run('Référencé par : ').bold = True
p.add_run('utilisateur.pays_id, donnee_collectee.pays_id.')

# --- utilisateur
add_heading(doc, 'Table utilisateur', level=2)
add_para(doc, 'Comptes des utilisateurs métier (distincts de auth_user qui gère l\'admin Django).')
add_table(doc,
    ['Colonne', 'Type', 'Contraintes', 'Description'],
    [
        ['id',                  'integer',      'PK, auto',                                    'Identifiant'],
        ['nom',                 'varchar(100)', 'NOT NULL',                                    'Nom de famille'],
        ['prenom',              'varchar(100)', 'NOT NULL',                                    'Prénom'],
        ['email',               'varchar(150)', 'NOT NULL, UNIQUE',                            'Identifiant de connexion'],
        ['mot_de_passe_hash',   'varchar(255)', 'NOT NULL',                                    'Hash du mot de passe'],
        ['institution',         'varchar(200)', 'NULL',                                        'Organisme de rattachement'],
        ['pays_id',             'integer',      'FK → pays(id)',                               'Pays affecté'],
        ['role',                'varchar(20)',  'NOT NULL, CHECK ∈ {collecteur, admin}',       'Profil applicatif'],
        ['actif',               'boolean',      'défaut true',                                 'Compte activé'],
        ['cree_le',             'timestamp',    'défaut now()',                                'Date de création'],
        ['derniere_connexion',  'timestamp',    'NULL',                                        'Dernière connexion'],
    ],
    col_widths=[3.5, 2.8, 4.5, 5.2])
p = doc.add_paragraph()
p.add_run('Référencé par : ').bold = True
p.add_run('donnee_collectee.utilisateur_id (saisie), donnee_collectee.valide_par (validation).')

# --- indicateur
doc.add_page_break()
add_heading(doc, 'Table indicateur', level=2)
add_para(doc, "Catalogue des 515 indicateurs issus des référentiels CISAT (climatique) "
              "et BASICSET (environnemental). Table la plus large : elle agrège métadonnées "
              "descriptives et références aux cadres internationaux.")
add_table(doc,
    ['Colonne', 'Type', 'Contraintes', 'Description'],
    [
        ['id',                  'integer',      'PK, auto',                                'Identifiant'],
        ['source',              'varchar(10)',  'NOT NULL, CHECK ∈ {CISAT, BASICSET}',     "Référentiel d'origine"],
        ['code',                'varchar(60)',  'UNIQUE, NULL',                            'Code court (ex. CISAT-1-A-01)'],
        ['numero_cisat',        'smallint',     'NULL',                                    'Numéro CISAT'],
        ['zone_cisat',          'varchar(50)',  'NULL',                                    'Zone thématique CISAT'],
        ['sujet',               'text',         'NULL',                                    'Sujet de l\'indicateur'],
        ['composante_basicset', 'varchar(200)', 'NULL',                                    'Composante BASICSET'],
        ['sous_composante',     'text',         'NULL',                                    'Sous-composante'],
        ['nom',                 'text',         'NOT NULL',                                'Libellé complet'],
        ['statistique',         'text',         'NULL',                                    'Statistique mesurée'],
        ['theme',               'varchar(100)', 'NULL',                                    'Thème'],
        ['etage',               'smallint',     'défaut 1, CHECK ∈ {1, 2, 3}',             'Niveau hiérarchique'],
        ['unite_mesure',        'varchar(150)', 'NULL',                                    'Unité (kg, %, °C, ...)'],
        ['accord_paris',        'varchar(100)', 'NULL',                                    'Article Accord de Paris'],
        ['pawp_katowice',       'text',         'NULL',                                    'Référence PAWP Katowice'],
        ['ref_fdes',            'varchar(200)', 'NULL',                                    'Référence FDES'],
        ['ref_odd',             'varchar(200)', 'NULL',                                    'Référence ODD'],
        ['ref_sendai',          'varchar(200)', 'NULL',                                    'Référence Cadre de Sendai'],
        ['ref_unece',           'varchar(300)', 'NULL',                                    'Référence UNECE'],
        ['ref_methodo',         'varchar(300)', 'NULL',                                    'Référence méthodologique'],
        ['sources_nat',         'text',         'NULL',                                    'Sources nationales'],
        ['institution_focale',  'text',         'NULL',                                    'Institution focale'],
        ['agregations',         'text',         'NULL',                                    'Agrégations possibles'],
        ['actif',               'boolean',      'défaut true',                             "Indicateur en cours d'usage"],
    ],
    col_widths=[3.5, 2.8, 4.5, 5.2])
p = doc.add_paragraph()
p.add_run('Index : ').bold = True
p.add_run('idx_ind_source (source), idx_ind_zone (zone_cisat), idx_ind_comp (composante_basicset), idx_ind_etage (etage).')
p = doc.add_paragraph()
p.add_run('Référencé par : ').bold = True
p.add_run('metadonnee_indicateur.indicateur_id (ON DELETE CASCADE), donnee_collectee.indicateur_id.')

# --- metadonnee
doc.add_page_break()
add_heading(doc, 'Table metadonnee_indicateur', level=2)
add_para(doc, "Métadonnées complémentaires en relation 1-1 avec indicateur. "
              "Isole les contenus longs sans alourdir les requêtes sur le catalogue.")
add_table(doc,
    ['Colonne', 'Type', 'Contraintes', 'Description'],
    [
        ['id',                       'integer',      'PK, auto',                                                          'Identifiant'],
        ['indicateur_id',            'integer',      'NOT NULL, UNIQUE, FK → indicateur(id) ON DELETE CASCADE',           'Indicateur associé'],
        ['definition',               'text',         'NULL',                                                              'Définition complète'],
        ['pertinence',               'text',         'NULL',                                                              'Pertinence / justification'],
        ['type_source_donnee',       'text',         'NULL',                                                              'Type de source (enquête, registre, ...)'],
        ['frequence_maj',            'text',         'NULL',                                                              'Fréquence de mise à jour'],
        ['categorie_mesure',         'varchar(150)', 'NULL',                                                              'Catégorie de mesure'],
        ['methode_calcul',           'text',         'NULL',                                                              'Formule / méthode de calcul'],
        ['agregations_potentielles', 'text',         'NULL',                                                              'Possibilités d\'agrégation'],
        ['cree_le',                  'timestamp',    'défaut now()',                                                      'Création'],
        ['modifie_le',               'timestamp',    'défaut now(), auto-MAJ via trigger',                                'Dernière modification'],
    ],
    col_widths=[4, 2.8, 5.5, 3.7])
p = doc.add_paragraph()
p.add_run('Trigger : ').bold = True
p.add_run('trg_meta_modifie_le (BEFORE UPDATE) → met à jour modifie_le.')

# --- donnee_collectee
doc.add_page_break()
add_heading(doc, 'Table donnee_collectee', level=2)
add_para(doc, "Table de faits : valeurs saisies par les collecteurs pour un triplet "
              "(indicateur, pays, année). Cœur opérationnel de l'application.")
add_table(doc,
    ['Colonne', 'Type', 'Contraintes', 'Description'],
    [
        ['id',               'integer',       'PK, auto',                                                                  'Identifiant'],
        ['indicateur_id',    'integer',       'NOT NULL, FK → indicateur(id)',                                             'Indicateur mesuré'],
        ['pays_id',          'integer',       'NOT NULL, FK → pays(id)',                                                   'Pays concerné'],
        ['utilisateur_id',   'integer',       'NOT NULL, FK → utilisateur(id)',                                            'Auteur de la saisie'],
        ['annee_reference',  'smallint',      'NOT NULL, CHECK 1990 ≤ x ≤ 2100',                                           'Année de référence'],
        ['valeur_numerique', 'numeric(20,4)', 'NULL',                                                                      'Valeur quantitative'],
        ['valeur_texte',     'text',          'NULL',                                                                      'Valeur qualitative'],
        ['unite_saisie',     'varchar(150)',  'NULL',                                                                      'Unité saisie'],
        ['source_donnee',    'text',          'NULL',                                                                      'Source'],
        ['methode_collecte', 'varchar(300)',  'NULL',                                                                      'Méthode de collecte'],
        ['commentaire',      'text',          'NULL',                                                                      'Note libre'],
        ['statut',           'varchar(20)',   "NOT NULL, défaut 'brouillon', CHECK ∈ {brouillon, soumis, valide, rejete}", 'État du workflow'],
        ['valide_par',       'integer',       'FK → utilisateur(id), NULL',                                                'Administrateur validant'],
        ['valide_le',        'timestamp',     'NULL',                                                                      'Date de validation'],
        ['motif_rejet',      'text',          'NULL',                                                                      'Raison du rejet'],
        ['cree_le',          'timestamp',     'défaut now()',                                                              'Création'],
        ['modifie_le',       'timestamp',     'défaut now(), auto-MAJ via trigger',                                        'Dernière modification'],
    ],
    col_widths=[3.5, 2.8, 5.5, 4.2])
p = doc.add_paragraph()
p.add_run('Contrainte UNIQUE : ').bold = True
p.add_run('(indicateur_id, pays_id, annee_reference) — une seule valeur par triplet.')
p = doc.add_paragraph()
p.add_run('Index : ').bold = True
p.add_run('idx_dc_indicateur, idx_dc_pays, idx_dc_annee, idx_dc_statut.')
p = doc.add_paragraph()
p.add_run('Trigger : ').bold = True
p.add_run('trg_donnee_modifie_le (BEFORE UPDATE) → met à jour modifie_le.')

# ---------------------------------------------------------------- §5
doc.add_page_break()
add_heading(doc, '5. Fonctions et triggers', level=1)
add_heading(doc, 'Fonction update_modifie_le()', level=2)

code = (
    "CREATE OR REPLACE FUNCTION public.update_modifie_le()\n"
    "RETURNS trigger\n"
    "LANGUAGE plpgsql\n"
    "AS $$\n"
    "BEGIN\n"
    "    NEW.modifie_le = NOW();\n"
    "    RETURN NEW;\n"
    "END;\n"
    "$$;"
)
p = doc.add_paragraph()
r = p.add_run(code)
r.font.name = 'Consolas'
r.font.size = Pt(10)

add_para(doc, 'Utilisée par les triggers BEFORE UPDATE sur :')
doc.add_paragraph('metadonnee_indicateur → trg_meta_modifie_le', style='List Bullet')
doc.add_paragraph('donnee_collectee → trg_donnee_modifie_le', style='List Bullet')

# ---------------------------------------------------------------- §6
add_heading(doc, '6. Workflow de validation des données', level=1)
workflow = (
    "[brouillon] ──soumission──► [soumis] ──admin valide──► [valide]\n"
    "                               │\n"
    "                               └──admin rejette──► [rejete]\n"
    "                                                     (motif_rejet renseigné)"
)
p = doc.add_paragraph()
r = p.add_run(workflow)
r.font.name = 'Consolas'
r.font.size = Pt(10)

doc.add_paragraph("Lors d'une validation : valide_par et valide_le sont renseignés.",
                  style='List Bullet')
doc.add_paragraph("Lors d'un rejet : motif_rejet doit être renseigné par l'administrateur.",
                  style='List Bullet')

# ---------------------------------------------------------------- §7
add_heading(doc, '7. Règles d\'intégrité clés', level=1)
rules = [
    "Unicité de saisie : impossible d'avoir deux valeurs pour le même (indicateur, pays, année).",
    "Suppression d'un indicateur : CASCADE sur metadonnee_indicateur, mais bloquée si des donnee_collectee y font référence.",
    "Suppression d'un utilisateur : bloquée s'il a saisi ou validé des données.",
    "Année plausible : bornée entre 1990 et 2100.",
    "Rôles : seulement collecteur ou admin.",
    "Sources d'indicateurs : seulement CISAT ou BASICSET.",
]
for i, rule in enumerate(rules, 1):
    doc.add_paragraph(f"{i}. {rule}", style='List Number')

# ---------------------------------------------------------------- §8
add_heading(doc, '8. Volumétrie actuelle (snapshot au 2026-05-20)', level=1)
add_table(doc,
    ['Table', 'Lignes'],
    [
        ['pays',                  '1'],
        ['utilisateur',           '1'],
        ['indicateur',            '515'],
        ['metadonnee_indicateur', '515'],
        ['donnee_collectee',      '3'],
    ],
    col_widths=[6, 3])

add_para(doc,
    "La base est encore en phase d'amorçage : catalogue d'indicateurs déjà chargé, "
    "collecte à démarrer.", italic=True)

# --------- save -----------------------------------------------------------
doc.save(OUT_PATH)
print(f"OK -> {OUT_PATH}")
